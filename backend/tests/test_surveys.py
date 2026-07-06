import pytest
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

import io
import json
import sqlite3
import uuid
from datetime import datetime, date
from unittest.mock import AsyncMock, patch
from fastapi.testclient import TestClient
from app.database import init_db
from app.services.auth_service import create_jwt


VALID_LLM_JSON = json.dumps({
    "sections": [
        {
            "title": "Bahagian A: Demografi",
            "questions": [
                {"question_text": "Jantina anda?", "question_type": "demographic",
                 "options": ["Lelaki", "Perempuan"], "likert_points": None},
            ],
        },
        {
            "title": "Bahagian B: Kepuasan Kerja",
            "questions": [
                {"question_text": "Saya berpuas hati dengan kerja saya.", "question_type": "likert",
                 "options": ["Sangat Tidak Setuju", "Tidak Setuju", "Tidak Pasti", "Setuju", "Sangat Setuju"],
                 "likert_points": 5},
                {"question_text": "Nyatakan cadangan penambahbaikan.", "question_type": "open",
                 "options": None, "likert_points": None},
            ],
        },
    ]
})


def make_token(user_id=None, email=None):
    uid = user_id or str(uuid.uuid4())
    em = email or f"test_{uuid.uuid4().hex[:6]}@test.com"
    return create_jwt({"user_id": uid, "email": em}), uid, em


@pytest.fixture
def client(tmp_path):
    db_path = str(tmp_path / "test.db")
    with patch("app.database._db_path", db_path):
        init_db(db_path)
        from app.main import app
        with TestClient(app) as c:
            yield c, db_path


def _reset_date():
    today = date.today()
    if today.month == 12:
        return date(today.year + 1, 1, 1).isoformat()
    return date(today.year, today.month + 1, 1).isoformat()


def _seed_user(db_path, user_id, email, kredit_subscription=50, kredit_topup=0):
    conn = sqlite3.connect(db_path)
    total = kredit_subscription + kredit_topup
    conn.execute(
        """INSERT OR IGNORE INTO users
           (id, email, tier, kredit_remaining, kredit_total, kredit_subscription, kredit_topup,
            tokens_used_internal, reset_date, created_at)
           VALUES (?, ?, 'free', ?, ?, ?, ?, 0, ?, ?)""",
        (user_id, email, total, total, kredit_subscription, kredit_topup,
         _reset_date(), datetime.utcnow().isoformat()),
    )
    conn.commit()
    conn.close()


def _seed_project(db_path, project_id, user_id):
    conn = sqlite3.connect(db_path)
    conn.execute(
        """INSERT INTO projects (id, user_id, title, research_mode, created_at)
           VALUES (?, ?, 'Test Project', 'general', ?)""",
        (project_id, user_id, datetime.utcnow().isoformat()),
    )
    conn.commit()
    conn.close()


def _seed_document(db_path, project_id):
    conn = sqlite3.connect(db_path)
    conn.execute(
        """INSERT INTO documents (id, project_id, filename, category, page_count, chunk_count, uploaded_at)
           VALUES (?, ?, 'proposal.pdf', 'proposal', 10, 5, ?)""",
        (str(uuid.uuid4()), project_id, datetime.utcnow().isoformat()),
    )
    conn.commit()
    conn.close()


def _setup(client_fixture, kredit_subscription=50, kredit_topup=0, with_doc=True):
    c, db_path = client_fixture
    token, uid, email = make_token()
    _seed_user(db_path, uid, email, kredit_subscription, kredit_topup)
    pid = str(uuid.uuid4())
    _seed_project(db_path, pid, uid)
    if with_doc:
        _seed_document(db_path, pid)
    headers = {"Authorization": f"Bearer {token}"}
    return c, db_path, uid, pid, headers


def _credits(db_path, uid):
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    row = conn.execute(
        "SELECT kredit_subscription, kredit_topup, kredit_remaining FROM users WHERE id=?", (uid,)
    ).fetchone()
    conn.close()
    return dict(row)


# ── CRUD: surveys ────────────────────────────────────────────────

def test_create_survey(client):
    c, db_path, uid, pid, headers = _setup(client)
    r = c.post(f"/projects/{pid}/surveys", json={}, headers=headers)
    assert r.status_code == 201
    body = r.json()
    assert body["title"] == "Soal Selidik"
    assert body["status"] == "draft"


def test_create_survey_other_user_project_404(client):
    c, db_path, uid, pid, headers = _setup(client)
    other_token, other_uid, other_email = make_token()
    _seed_user(db_path, other_uid, other_email)
    r = c.post(f"/projects/{pid}/surveys", json={}, headers={"Authorization": f"Bearer {other_token}"})
    assert r.status_code == 404


def test_list_surveys(client):
    c, db_path, uid, pid, headers = _setup(client)
    c.post(f"/projects/{pid}/surveys", json={"title": "Survey Satu"}, headers=headers)
    c.post(f"/projects/{pid}/surveys", json={"title": "Survey Dua"}, headers=headers)
    r = c.get(f"/projects/{pid}/surveys", headers=headers)
    assert r.status_code == 200
    titles = {s["title"] for s in r.json()}
    assert titles == {"Survey Satu", "Survey Dua"}


def test_rename_survey(client):
    c, db_path, uid, pid, headers = _setup(client)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    r = c.patch(f"/surveys/{sid}", json={"title": "Instrumen Kajian Saya"}, headers=headers)
    assert r.status_code == 200
    assert c.get(f"/surveys/{sid}", headers=headers).json()["title"] == "Instrumen Kajian Saya"


def test_survey_ownership_other_user_gets_404(client):
    c, db_path, uid, pid, headers = _setup(client)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    other_token, other_uid, other_email = make_token()
    _seed_user(db_path, other_uid, other_email)
    other_headers = {"Authorization": f"Bearer {other_token}"}
    assert c.get(f"/surveys/{sid}", headers=other_headers).status_code == 404
    assert c.patch(f"/surveys/{sid}", json={"title": "hijack"}, headers=other_headers).status_code == 404
    assert c.delete(f"/surveys/{sid}", headers=other_headers).status_code == 404


def test_delete_survey_cascades_sections_and_questions(client):
    c, db_path, uid, pid, headers = _setup(client)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    sec = c.post(f"/surveys/{sid}/sections", json={"title": "Bahagian A"}, headers=headers).json()
    c.post(f"/sections/{sec['id']}/questions",
           json={"question_text": "Umur anda?", "question_type": "demographic", "options": ["<25", "25-40", ">40"]},
           headers=headers)
    r = c.delete(f"/surveys/{sid}", headers=headers)
    assert r.status_code == 204
    conn = sqlite3.connect(db_path)
    conn.execute("PRAGMA foreign_keys = ON")
    assert conn.execute("SELECT COUNT(*) FROM survey_sections WHERE survey_id=?", (sid,)).fetchone()[0] == 0
    assert conn.execute("SELECT COUNT(*) FROM survey_questions").fetchone()[0] == 0
    conn.close()


# ── CRUD: sections & questions ───────────────────────────────────

def test_section_crud(client):
    c, db_path, uid, pid, headers = _setup(client)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    r = c.post(f"/surveys/{sid}/sections", json={"title": "Bahagian A: Demografi"}, headers=headers)
    assert r.status_code == 201
    sec_id = r.json()["id"]
    assert r.json()["position"] == 0
    r2 = c.patch(f"/sections/{sec_id}", json={"title": "Bahagian A (Edited)", "position": 5}, headers=headers)
    assert r2.status_code == 200
    assert r2.json()["title"] == "Bahagian A (Edited)"
    assert r2.json()["position"] == 5
    assert c.delete(f"/sections/{sec_id}", headers=headers).status_code == 204
    assert c.get(f"/surveys/{sid}", headers=headers).json()["sections"] == []


def test_question_crud(client):
    c, db_path, uid, pid, headers = _setup(client)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    sec_id = c.post(f"/surveys/{sid}/sections", json={"title": "B"}, headers=headers).json()["id"]
    r = c.post(f"/sections/{sec_id}/questions",
               json={"question_text": "Saya suka kerja saya.", "question_type": "likert",
                     "options": ["STS", "TS", "TP", "S", "SS"], "likert_points": 5},
               headers=headers)
    assert r.status_code == 201
    q_id = r.json()["id"]
    r2 = c.patch(f"/questions/{q_id}",
                 json={"question_text": "Saya tidak suka kerja saya.", "is_reversed": True},
                 headers=headers)
    assert r2.status_code == 200
    assert r2.json()["is_reversed"] is True
    full = c.get(f"/surveys/{sid}", headers=headers).json()
    q = full["sections"][0]["questions"][0]
    assert q["question_text"] == "Saya tidak suka kerja saya."
    assert q["is_reversed"] is True
    assert q["options"] == ["STS", "TS", "TP", "S", "SS"]
    assert c.delete(f"/questions/{q_id}", headers=headers).status_code == 204


def test_position_ordering(client):
    c, db_path, uid, pid, headers = _setup(client)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    sec_id = c.post(f"/surveys/{sid}/sections", json={"title": "B"}, headers=headers).json()["id"]
    qa = c.post(f"/sections/{sec_id}/questions", json={"question_text": "Soalan A"}, headers=headers).json()
    qb = c.post(f"/sections/{sec_id}/questions", json={"question_text": "Soalan B"}, headers=headers).json()
    assert (qa["position"], qb["position"]) == (0, 1)
    # swap positions
    c.patch(f"/questions/{qa['id']}", json={"position": 1}, headers=headers)
    c.patch(f"/questions/{qb['id']}", json={"position": 0}, headers=headers)
    full = c.get(f"/surveys/{sid}", headers=headers).json()
    texts = [q["question_text"] for q in full["sections"][0]["questions"]]
    assert texts == ["Soalan B", "Soalan A"]


def test_invalid_question_type_rejected(client):
    c, db_path, uid, pid, headers = _setup(client)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    sec_id = c.post(f"/surveys/{sid}/sections", json={"title": "B"}, headers=headers).json()["id"]
    r = c.post(f"/sections/{sec_id}/questions",
               json={"question_text": "X", "question_type": "ranking"}, headers=headers)
    assert r.status_code == 400


# ── AI Generation ────────────────────────────────────────────────

def test_generate_full_success_deducts_10_credits(client):
    c, db_path, uid, pid, headers = _setup(client, kredit_subscription=50)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    with patch("app.services.survey_generator.get_project_context", new_callable=AsyncMock, return_value="konteks kajian"), \
         patch("app.services.survey_generator.call_deepseek_raw", new_callable=AsyncMock, return_value=VALID_LLM_JSON):
        r = c.post(f"/surveys/{sid}/generate", json={"scope": "full"}, headers=headers)
    assert r.status_code == 200
    body = r.json()
    assert len(body["sections"]) == 2
    assert body["sections"][0]["title"] == "Bahagian A: Demografi"
    assert body["sections"][1]["questions"][0]["question_type"] == "likert"
    assert body["sections"][1]["questions"][0]["likert_points"] == 5
    assert body["kredit_used"] == 10
    credits = _credits(db_path, uid)
    assert credits["kredit_subscription"] == 40
    assert credits["kredit_remaining"] == 40


def test_generate_deduction_order_subscription_then_topup(client):
    c, db_path, uid, pid, headers = _setup(client, kredit_subscription=5, kredit_topup=20)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    with patch("app.services.survey_generator.get_project_context", new_callable=AsyncMock, return_value="konteks"), \
         patch("app.services.survey_generator.call_deepseek_raw", new_callable=AsyncMock, return_value=VALID_LLM_JSON):
        r = c.post(f"/surveys/{sid}/generate", json={"scope": "full"}, headers=headers)
    assert r.status_code == 200
    credits = _credits(db_path, uid)
    assert credits["kredit_subscription"] == 0
    assert credits["kredit_topup"] == 15
    assert credits["kredit_remaining"] == 15


def test_generate_no_documents_400_no_deduction(client):
    c, db_path, uid, pid, headers = _setup(client, with_doc=False)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    r = c.post(f"/surveys/{sid}/generate", json={"scope": "full"}, headers=headers)
    assert r.status_code == 400
    assert "Muat naik" in r.json()["detail"]
    assert _credits(db_path, uid)["kredit_remaining"] == 50


def test_generate_invalid_json_twice_error_no_deduction(client):
    c, db_path, uid, pid, headers = _setup(client)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    mock_llm = AsyncMock(return_value="ini bukan JSON langsung")
    with patch("app.services.survey_generator.get_project_context", new_callable=AsyncMock, return_value="konteks"), \
         patch("app.services.survey_generator.call_deepseek_raw", mock_llm):
        r = c.post(f"/surveys/{sid}/generate", json={"scope": "full"}, headers=headers)
    assert r.status_code == 502
    assert mock_llm.call_count == 2  # retry SEKALI
    assert _credits(db_path, uid)["kredit_remaining"] == 50
    # DB tidak berubah
    assert c.get(f"/surveys/{sid}", headers=headers).json()["sections"] == []


def test_generate_insufficient_credits_402(client):
    c, db_path, uid, pid, headers = _setup(client, kredit_subscription=4, kredit_topup=0)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    r = c.post(f"/surveys/{sid}/generate", json={"scope": "full"}, headers=headers)
    assert r.status_code == 402
    assert _credits(db_path, uid)["kredit_remaining"] == 4
    assert c.get(f"/surveys/{sid}", headers=headers).json()["sections"] == []


def test_generate_section_scope_costs_3_and_appends(client):
    c, db_path, uid, pid, headers = _setup(client, kredit_subscription=50)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    c.post(f"/surveys/{sid}/sections", json={"title": "Bahagian Sedia Ada"}, headers=headers)
    section_json = json.dumps({"sections": [{"title": "Bahagian C: Motivasi", "questions": [
        {"question_text": "Saya bermotivasi.", "question_type": "likert",
         "options": ["STS", "TS", "TP", "S", "SS"], "likert_points": 5}]}]})
    with patch("app.services.survey_generator.get_project_context", new_callable=AsyncMock, return_value="konteks"), \
         patch("app.services.survey_generator.call_deepseek_raw", new_callable=AsyncMock, return_value=section_json):
        r = c.post(f"/surveys/{sid}/generate",
                   json={"scope": "section", "instruction": "jana bahagian motivasi"}, headers=headers)
    assert r.status_code == 200
    body = r.json()
    assert body["kredit_used"] == 3
    # APPEND, bukan replace
    titles = [s["title"] for s in body["sections"]]
    assert titles == ["Bahagian Sedia Ada", "Bahagian C: Motivasi"]
    assert _credits(db_path, uid)["kredit_remaining"] == 47


def test_generate_full_replaces_existing_content(client):
    c, db_path, uid, pid, headers = _setup(client)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    c.post(f"/surveys/{sid}/sections", json={"title": "Lama"}, headers=headers)
    with patch("app.services.survey_generator.get_project_context", new_callable=AsyncMock, return_value="konteks"), \
         patch("app.services.survey_generator.call_deepseek_raw", new_callable=AsyncMock, return_value=VALID_LLM_JSON):
        r = c.post(f"/surveys/{sid}/generate", json={"scope": "full"}, headers=headers)
    titles = [s["title"] for s in r.json()["sections"]]
    assert "Lama" not in titles
    assert len(titles) == 2


def test_generate_markdown_fenced_json_parsed(client):
    c, db_path, uid, pid, headers = _setup(client)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    fenced = f"```json\n{VALID_LLM_JSON}\n```"
    with patch("app.services.survey_generator.get_project_context", new_callable=AsyncMock, return_value="konteks"), \
         patch("app.services.survey_generator.call_deepseek_raw", new_callable=AsyncMock, return_value=fenced):
        r = c.post(f"/surveys/{sid}/generate", json={"scope": "full"}, headers=headers)
    assert r.status_code == 200
    assert len(r.json()["sections"]) == 2


# ── Export ───────────────────────────────────────────────────────

def test_export_docx_roundtrip(client):
    from docx import Document
    c, db_path, uid, pid, headers = _setup(client)
    sid = c.post(f"/projects/{pid}/surveys", json={"title": "Instrumen Ujian"}, headers=headers).json()["id"]
    sec_id = c.post(f"/surveys/{sid}/sections", json={"title": "Bahagian A: Demografi"}, headers=headers).json()["id"]
    c.post(f"/sections/{sec_id}/questions",
           json={"question_text": "Jantina anda?", "question_type": "mcq", "options": ["Lelaki", "Perempuan"]},
           headers=headers)
    c.post(f"/sections/{sec_id}/questions",
           json={"question_text": "Saya berpuas hati.", "question_type": "likert",
                 "options": ["STS", "TS", "TP", "S", "SS"], "likert_points": 5},
           headers=headers)
    c.post(f"/sections/{sec_id}/questions",
           json={"question_text": "Cadangan anda?", "question_type": "open"},
           headers=headers)

    r = c.get(f"/surveys/{sid}/export/docx", headers=headers)
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    doc = Document(io.BytesIO(r.content))  # round-trip: mesti boleh dibuka
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Instrumen Ujian" in all_text
    assert "Bahagian A: Demografi" in all_text
    assert "1. Jantina anda?" in all_text
    assert "☐ Lelaki" in all_text
    assert len(doc.tables) == 1  # likert → jadual skala
    assert doc.tables[0].rows[0].cells[0].text == "STS"


def test_export_ownership_other_user_404(client):
    c, db_path, uid, pid, headers = _setup(client)
    sid = c.post(f"/projects/{pid}/surveys", json={}, headers=headers).json()["id"]
    other_token, other_uid, other_email = make_token()
    _seed_user(db_path, other_uid, other_email)
    r = c.get(f"/surveys/{sid}/export/docx", headers={"Authorization": f"Bearer {other_token}"})
    assert r.status_code == 404

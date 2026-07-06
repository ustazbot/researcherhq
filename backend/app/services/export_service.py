import asyncio
import io
import re
import uuid
from typing import Dict
from bs4 import BeautifulSoup

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def _html_to_docx_paragraphs(doc, html_content: str):
    """Convert TipTap HTML content to python-docx paragraphs with formatting."""
    html_content = re.sub(r'\[\[cite:\d+\]\]', '', html_content)
    html_content = html_content.strip()
    if not html_content:
        return

    soup = BeautifulSoup(html_content, 'html.parser')

    for el in soup.children:
        if not hasattr(el, 'name') or not el.name:
            continue
        tag = el.name

        if tag in ('h1', 'h2', 'h3', 'h4'):
            level = int(tag[1])
            text = el.get_text(strip=True)
            if text:
                doc.add_heading(text, level=level)

        elif tag == 'p':
            text = el.get_text(strip=True)
            if not text:
                continue
            para = doc.add_paragraph()
            for child in el.children:
                if not hasattr(child, 'name') or child.name is None:
                    raw = str(child)
                    if raw:
                        para.add_run(raw)
                elif child.name == 'strong':
                    run = para.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em':
                    run = para.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'u':
                    run = para.add_run(child.get_text())
                    run.underline = True
                elif child.name == 'mark':
                    para.add_run(child.get_text())
                else:
                    para.add_run(child.get_text())

        elif tag == 'ul':
            for li in el.find_all('li', recursive=False):
                text = li.get_text(separator=' ', strip=True)
                if text:
                    doc.add_paragraph(text, style='List Bullet')

        elif tag == 'ol':
            for li in el.find_all('li', recursive=False):
                text = li.get_text(separator=' ', strip=True)
                if text:
                    doc.add_paragraph(text, style='List Number')

        elif tag == 'blockquote':
            text = el.get_text(separator=' ', strip=True)
            if text:
                para = doc.add_paragraph(text)
                para.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else 'Normal'


# job_id → {"status": "pending"|"done"|"error", "bytes": bytes|None, "filename": str}
_jobs: Dict[str, dict] = {}
_queue: asyncio.Queue = None


def get_queue() -> asyncio.Queue:
    global _queue
    if _queue is None:
        _queue = asyncio.Queue()
    return _queue


async def start_export_worker():
    global _queue
    _queue = asyncio.Queue()
    q = _queue
    while True:
        item = await q.get()
        job_id = item["job_id"]
        try:
            if item["type"] == "chapter":
                result = await asyncio.get_event_loop().run_in_executor(
                    None, _build_docx, item["chapter_title"], item["content"]
                )
                _jobs[job_id] = {"status": "done", "bytes": result, "filename": f"{item['chapter_title']}.docx"}
            elif item["type"] == "thesis":
                result = await asyncio.get_event_loop().run_in_executor(
                    None, _build_thesis_docx,
                    item["project_title"], item["chapters"],
                    item["bibliography"], item["citation_style"]
                )
                safe_title = item["project_title"][:50].replace("/", "-")
                _jobs[job_id] = {"status": "done", "bytes": result, "filename": f"{safe_title}.docx"}
        except Exception as e:
            _jobs[job_id] = {"status": "error", "bytes": None, "filename": "", "error": str(e)}
        finally:
            q.task_done()


def _build_docx(chapter_title: str, content: str) -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx tidak dipasang.")
    doc = Document()
    doc.add_heading(chapter_title, level=1)
    _html_to_docx_paragraphs(doc, content)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _insert_toc(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run()
    run = paragraph.runs[0]
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _build_thesis_docx(project_title: str, chapters: list, bibliography: list, citation_style: str) -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx tidak dipasang.")

    doc = Document()

    # Cover page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(project_title)
    run.font.size = Pt(24)
    run.bold = True
    doc.add_page_break()

    # TOC field (auto-builds when user opens in Word)
    _insert_toc(doc)
    doc.add_page_break()

    # Chapters sorted: front_matter → chapter → appendix
    ORDER = {"front_matter": 0, "chapter": 1, "appendix": 2}
    sorted_chapters = sorted(chapters, key=lambda c: (ORDER.get(c["section_type"], 1), c["chapter_order"]))

    for chap in sorted_chapters:
        if not chap["content"].strip():
            continue
        doc.add_heading(chap["title"], level=1)
        _html_to_docx_paragraphs(doc, chap["content"])
        doc.add_page_break()

    # Bibliography
    if bibliography:
        doc.add_heading("Senarai Rujukan", level=1)
        seen_files = {}
        for src in bibliography:
            fname = src["filename"]
            if fname not in seen_files:
                seen_files[fname] = []
            seen_files[fname].append(src["page_number"])

        for fname, pages in seen_files.items():
            pages_str = ", ".join(f"ms. {p}" for p in sorted(set(pages)))
            doc.add_paragraph(f"{fname} ({pages_str})", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def enqueue_export(chapter_title: str, content: str) -> str:
    job_id = str(uuid.uuid4())
    _jobs[job_id] = {"status": "pending", "bytes": None, "filename": ""}
    get_queue().put_nowait({
        "type": "chapter",
        "job_id": job_id,
        "chapter_title": chapter_title,
        "content": content,
    })
    return job_id


def enqueue_thesis_compile(project_title: str, chapters: list, bibliography: list, citation_style: str) -> str:
    job_id = str(uuid.uuid4())
    _jobs[job_id] = {"status": "pending", "bytes": None, "filename": ""}
    get_queue().put_nowait({
        "type": "thesis",
        "job_id": job_id,
        "project_title": project_title,
        "chapters": chapters,
        "bibliography": bibliography,
        "citation_style": citation_style,
    })
    return job_id


def get_job(job_id: str) -> dict | None:
    return _jobs.get(job_id)


def build_survey_docx(survey_title: str, sections: list) -> bytes:
    """Task 36A: survey instrument → .docx, in-memory, zero disk write.
    sections = [{title, questions: [{question_text, question_type, options(list|None), likert_points}]}]"""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx tidak dipasang.")
    doc = Document()
    doc.add_heading(survey_title, level=1)

    for section in sections:
        doc.add_heading(section["title"], level=2)
        q_num = 0
        for q in section["questions"]:
            q_num += 1
            doc.add_paragraph(f"{q_num}. {q['question_text']}")
            qtype = q["question_type"]
            options = q.get("options") or []
            if qtype == "likert" and options:
                table = doc.add_table(rows=2, cols=len(options))
                table.style = "Table Grid"
                for i, label in enumerate(options):
                    table.rows[0].cells[i].text = str(label)
                    table.rows[1].cells[i].text = str(i + 1)
            elif qtype in ("mcq", "demographic") and options:
                for opt in options:
                    doc.add_paragraph(f"☐ {opt}", style="List Bullet")
            else:  # open-ended: blank writing space
                doc.add_paragraph("_" * 60)
                doc.add_paragraph("_" * 60)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_apa_docx(title: str, apa_tables: list) -> bytes:
    """Task 36C-1: render APA-style analysis tables to .docx, in-memory.
    apa_tables = [{title, columns, rows, note}]."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx tidak dipasang.")
    doc = Document()
    doc.add_heading(title, level=1)
    for t in apa_tables:
        # APA table titles are italic
        p = doc.add_paragraph()
        run = p.add_run(t.get("title", ""))
        run.italic = True
        run.bold = True
        cols = t.get("columns", [])
        rows = t.get("rows", [])
        if cols:
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Table Grid"
            for i, c in enumerate(cols):
                table.rows[0].cells[i].text = str(c)
            for r in rows:
                cells = table.add_row().cells
                for i, v in enumerate(r):
                    cells[i].text = "" if v is None else str(v)
        note = t.get("note")
        if note:
            np_ = doc.add_paragraph()
            nr = np_.add_run(f"Note. {note}")
            nr.italic = True
        doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
